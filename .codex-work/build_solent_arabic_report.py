from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Yazan\Desktop\Projects\Solent\app")
OUTPUT_DIR = ROOT / "outputs" / "client-report"
OUTPUT_FILE = OUTPUT_DIR / "Solent_Arabic_System_Overview.docx"
LOGO = ROOT / "public" / "images" / "brands" / "solent" / "solent_h_png.png"

NAVY = "0F172A"
TEAL = "0F766E"
TEAL_LIGHT = "E6F4F1"
BLUE_LIGHT = "EFF6FF"
AMBER_LIGHT = "FFF7E6"
GRAY_LIGHT = "F4F6F8"
GRAY = "52606D"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), "Arial")
    if r_pr.find(qn("w:rtl")) is None:
        r_pr.append(OxmlElement("w:rtl"))
    if r_pr.find(qn("w:lang")) is None:
        lang = OxmlElement("w:lang")
        lang.set(qn("w:val"), "ar-JO")
        lang.set(qn("w:bidi"), "ar-JO")
        r_pr.append(lang)


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._element.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_paragraph_border(paragraph, color=TEAL, size="18", space="7", side="bottom"):
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=180, bottom=120, end=180):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, TEAL, 16, 8),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 12, NAVY, 9, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(section, cover=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1.15 if cover else 0.85)
    section.bottom_margin = Inches(0.8)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)


def create_numbering(doc, fmt="bullet"):
    numbering = doc.part.numbering_part.element
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abs_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "right")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    p_pr.append(OxmlElement("w:bidi"))
    ind = OxmlElement("w:ind")
    ind.set(qn("w:right"), "400")
    ind.set(qn("w:hanging"), "240")
    p_pr.append(ind)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), "Arial")
    r_pr.append(fonts)
    r_pr.append(OxmlElement("w:rtl"))
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_num_pr(paragraph, num_id):
    p_pr = paragraph._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_ref)
    p_pr.append(num_pr)


def add_body(doc, text, bold_lead=None, color=NAVY, size=11, after=6, keep=False):
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=size, bold=True, color=TEAL)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=size, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_rtl(p)
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=TEAL if level < 3 else NAVY)
    if level == 1:
        set_paragraph_border(p, color="B9DDD6", size="8", space="4")
    return p


def add_list(doc, items, numbered=False, compact=False):
    num_id = create_numbering(doc, "decimal" if numbered else "bullet")
    for item in items:
        p = doc.add_paragraph()
        set_paragraph_rtl(p)
        add_num_pr(p, num_id)
        p.paragraph_format.space_after = Pt(3 if compact else 5)
        p.paragraph_format.line_spacing = 1.06
        p.paragraph_format.keep_together = True
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            set_run_font(r1, size=10.8, bold=True, color=NAVY)
            r2 = p.add_run(rest)
            set_run_font(r2, size=10.8, color=NAVY)
        else:
            run = p.add_run(item)
            set_run_font(run, size=10.8, color=NAVY)


def add_callout(doc, title, text, fill=TEAL_LIGHT, title_color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, start=220, bottom=140, end=220)
    p = cell.paragraphs[0]
    set_paragraph_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title)
    set_run_font(r1, size=11, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    set_paragraph_rtl(p2)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.06
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.4, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run("صفحة ")
    set_run_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_running_label(doc):
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Solent  |  نظرة شاملة على النظام")
    set_run_font(run, size=9, bold=True, color=TEAL)
    set_paragraph_border(p, color="D7E3E0", size="4", space="4", side="bottom")


def add_page_break(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_page(section)
    add_running_label(doc)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    configure_page(doc.sections[0], cover=True)

    props = doc.core_properties
    props.title = "نظرة شاملة على نظام Solent"
    props.subject = "تعريف عملي بنظام إدارة مختبر الأسنان واقتراح إضافة مخزون خفيف"
    props.author = "Solent"
    props.keywords = "Solent, dental lab, نظام مختبر أسنان, مخزون"

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(2.35))
        logo_shape._inline.docPr.set("title", "Solent")
        logo_shape._inline.docPr.set("descr", "شعار Solent")
    p.paragraph_format.space_after = Pt(52)

    title = doc.add_paragraph()
    set_paragraph_rtl(title, WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("نظرة شاملة على نظام Solent")
    set_run_font(r, size=29, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    set_paragraph_rtl(sub, WD_ALIGN_PARAGRAPH.CENTER)
    sub.paragraph_format.space_after = Pt(24)
    r = sub.add_run("من تسجيل الحالة إلى الإنتاج والتسليم والحسابات")
    set_run_font(r, size=16, color=TEAL)
    set_paragraph_border(sub, color=TEAL, size="14", space="12")

    note = doc.add_paragraph()
    set_paragraph_rtl(note, WD_ALIGN_PARAGRAPH.CENTER)
    note.paragraph_format.space_after = Pt(8)
    r = note.add_run("وثيقة تعريفية مستقلة للمراجعة واتخاذ القرار")
    set_run_font(r, size=12, bold=True, color=NAVY)

    date = doc.add_paragraph()
    set_paragraph_rtl(date, WD_ALIGN_PARAGRAPH.CENTER)
    date.paragraph_format.space_after = Pt(90)
    r = date.add_run("أغسطس 2026")
    set_run_font(r, size=10.5, color=GRAY)

    scope = doc.add_paragraph()
    set_paragraph_rtl(scope, WD_ALIGN_PARAGRAPH.CENTER)
    r = scope.add_run("هذه الوثيقة تصف النظام كما هو حالياً، وتفصل بين الميزات الجاهزة والأفكار المقترحة للتطوير.")
    set_run_font(r, size=9.5, color=GRAY)

    # Main content section
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_page(section)
    section.footer.is_linked_to_previous = False
    add_page_number(section.footer.paragraphs[0])
    add_running_label(doc)

    add_heading(doc, "1. الصورة العامة", 1)
    add_body(
        doc,
        "Solent هو نظام لإدارة العمل اليومي داخل مختبر الأسنان. يربط سجل العميل بالحالة، ثم يمرر الأعمال بين مراحل الإنتاج، ويتابع الجودة والتسليم، وبعدها يربط النتيجة بالفواتير والدفعات والتقارير. الفكرة ببساطة: تكون معلومات الحالة في مكان واحد بدل ما تتوزع بين الورق وواتساب وملفات منفصلة.",
        keep=True,
    )
    add_callout(
        doc,
        "ما الذي يحاول النظام حله؟",
        "تقليل ضياع المعلومات، توضيح من يعمل على كل حالة، معرفة المرحلة الحالية وموعد التسليم، وربط العمل المنجز بحساب العميل بشكل قابل للمراجعة.",
    )

    add_heading(doc, "رحلة الحالة داخل النظام", 2)
    add_list(
        doc,
        [
            ("استلام الحالة: ", "إنشاء رقم مرجعي وتسجيل العميل والمريض وموعد التسليم."),
            ("تعريف العمل: ", "إضافة عمل واحد أو أكثر مع الوحدات والأسنان أو الفك والمادة واللون والتفاصيل الفنية."),
            ("الإنتاج: ", "نقل كل عمل في المسار المناسب له حسب المادة، وليس في مسار ثابت للجميع."),
            ("المتابعة: ", "عرض الأعمال المنتظرة والنشطة وتسجيل من استلمها ومن أنجزها."),
            ("الجودة: ", "توثيق QC وأي Reject أو Repeat أو Modify أو Redo مع السبب."),
            ("التسليم: ", "متابعة الموعد، تجهيز جدول التسليم، وتسليم الحالة عبر الموظف أو السائق."),
            ("الحساب: ", "إصدار الفاتورة وتسجيل الدفعات وإظهار كشف حساب العميل."),
        ],
        numbered=True,
        compact=True,
    )

    add_page_break(doc)
    add_heading(doc, "2. سجل العميل أو الطبيب", 1)
    add_body(
        doc,
        "سجل العميل هو نقطة الربط بين الشغل الفني والحسابات. عملياً، الإدارة تقدر تفتح ملف الطبيب وتشوف بياناته وحالاته وتعاملاته بدون البحث في أكثر من مكان.",
        keep=True,
    )
    add_list(
        doc,
        [
            ("البيانات الأساسية: ", "الاسم، الهاتف الشخصي، هاتف العيادة، العنوان، وحالة الحساب فعّال أو متوقف."),
            ("استمرارية السجل: ", "إيقاف العميل لا يحذف تاريخه؛ الحالات والحسابات السابقة تبقى محفوظة ويمكن إعادة تفعيله."),
            ("التسعير الخاص: ", "إمكانية وضع خصم ثابت أو نسبة على مادة محددة لهذا الطبيب، مع خصم إضافي على حالة بعينها عند الحاجة."),
            ("الرصيد: ", "عرض الرصيد الحالي أو حتى تاريخ محدد، وإدخال رصيد افتتاحي عند الانتقال من نظام سابق."),
            ("الوصول السريع: ", "من ملف الطبيب يمكن الانتقال للحالات والفواتير والدفعات وكشف الحساب."),
        ],
    )
    add_callout(
        doc,
        "مهم قبل العرض على العميل",
        "توجد بداية لبوابة إلكترونية تعرض الرصيد والفواتير والدفعات، لكنها ليست جاهزة كخدمة مستقلة للعميل حالياً؛ تحتاج إكمال الدخول واختبار الصلاحيات والتجربة كاملة.",
        fill=AMBER_LIGHT,
        title_color="9A6700",
    )

    add_heading(doc, "3. ملف الحالة", 1)
    add_body(doc, "الحالة الواحدة يمكن أن تحتوي أكثر من عمل، وهذا مناسب للحالات المركبة التي تجمع مواد أو أنواع شغل مختلفة.")
    add_list(
        doc,
        [
            "رقم مرجعي تلقائي، الطبيب، اسم المريض، نوع الطبعة، وتاريخ التسليم المتوقع.",
            "أعمال متعددة داخل الحالة مع عدد الوحدات والأسنان أو الفك، Job Type، المادة، اللون أو الـStyle، وتفاصيل Implant أو Abutment عند الحاجة.",
            "ملاحظات متعددة، صور ومرفقات مرتبطة بالحالة، Tags للتنظيم، وخصم خاص بالحالة مع سبب واضح.",
            "بحث وفلترة حسب الطبيب والتاريخ والمريض، مع قفل الحالة والتحكم بالتعديل حسب الصلاحية.",
            "حذف مرن وإعادة استرجاع، وطباعة Label عادي أو مختصر، وطباعة فاتورة بعد اكتمال العمل.",
        ],
        compact=True,
    )

    add_page_break(doc)
    add_heading(doc, "4. التشغيل ومراحل الإنتاج", 1)
    add_body(
        doc,
        "كل مادة تحدد المراحل التي تحتاجها. لذلك ممكن أن يمر عمل عبر Design وMilling ثم Sintering وFinishing وQC، بينما عمل آخر يذهب إلى 3D Printing أو Pressing أو Metal Work. النظام يتجاوز المراحل غير المطلوبة تلقائياً.",
        keep=True,
    )
    add_list(
        doc,
        [
            ("قوائم واضحة: ", "لكل مرحلة أعمال منتظرة وأعمال نشطة."),
            ("استلام وتوزيع: ", "الموظف يستطيع أخذ العمل، والإدارة تستطيع التعيين أو إعادة التعيين حسب الصلاحيات."),
            ("إكمال المرحلة: ", "عند الإنهاء ينتقل العمل إلى المرحلة المناسبة التالية، مع حفظ سجل العملية والمستخدم."),
            ("إدارة الاستثناء: ", "يمكن إرجاع العمل للانتظار أو تجاوز خطوة عند وجود سبب وصلاحية واضحة."),
            ("متابعة مباشرة: ", "Operations Dashboard وLive Monitor يساعدان في معرفة ضغط الشغل ومكان كل حالة."),
        ],
    )
    add_callout(
        doc,
        "حدود هذه النقطة",
        "النظام يدير انتقال الأعمال بين المراحل، لكنه لا يدير الماكينات أو الأجهزة حالياً. أي شاشة قديمة للـDevices أو Builds لا تُعتبر ميزة عاملة في النسخة الحالية.",
        fill=AMBER_LIGHT,
        title_color="9A6700",
    )

    add_heading(doc, "5. الجودة وإعادة العمل", 1)
    add_body(doc, "مش كل مشكلة هي نفس الشيء؛ النظام يفصل بين أنواع الإعادة حتى تكون الأرقام عادلة ويسهل معرفة السبب الحقيقي.")
    add_list(
        doc,
        [
            ("Reject: ", "رفض العمل في مرحلة معينة مع السبب والتوضيح."),
            ("Repeat: ", "إعادة تصنيع العمل وربطه بالأصل."),
            ("Modify: ", "طلب تعديل على العمل الموجود."),
            ("Redo: ", "إعادة كاملة مع الاحتفاظ بعلاقة واضحة مع الحالة السابقة."),
            ("التوثيق: ", "حفظ السبب، الموظف، الوقت، الملاحظات والـTags، ثم إظهارها في تقارير الجودة."),
        ],
        compact=True,
    )

    add_page_break(doc)
    add_heading(doc, "6. التوصيل وجدولة المواعيد", 1)
    add_body(doc, "جدول التسليم يعطي الفريق قائمة بالحالات غير المسلّمة ضمن فترة، ويبيّن الطبيب والمريض والوقت والوحدات والحالة الحالية، مع تمييز المتأخر واليوم والغد والأسبوع.")
    add_list(
        doc,
        [
            "تعديل موعد التسليم مع تسجيل المستخدم وإضافة ملاحظة على الحالة.",
            "طباعة جدول التسليم لتنسيق الجولة أو العمل داخل المختبر.",
            "قائمة للسائق: منتظر، نشط، وتم التسليم؛ يستلم السائق الحالة ثم يؤكد إنجاز التوصيل.",
            "ارتباط التسليم بإنهاء الحالة والفاتورة، مع بقاء سجل الحركة للمراجعة.",
        ],
    )

    add_heading(doc, "7. الفواتير والدفعات وكشف الحساب", 1)
    add_body(doc, "الحسابات في Solent مرتبطة بالشغل الفعلي، وليست دفتر حساب منفصل. قيمة الفاتورة تبدأ من عدد الوحدات وسعر المادة، ثم تدخل خصومات الطبيب وخصم الحالة.")
    add_list(
        doc,
        [
            ("الفواتير: ", "إنشاء فاتورة للحالة وإضافتها إلى حساب الطبيب عند اكتمالها، مع تعديل الرصيد إذا تغيرت حالة مفوترة."),
            ("الدفعات: ", "نقدي، شيك، حوالة بنكية أو CliQ، مع البنك أو رقم الشيك والملاحظات والموظف والتاريخ."),
            ("الاستلام الداخلي: ", "المحاسبة تستطيع تأكيد استلام الدفعة من الموظف الذي حصّلها."),
            ("التسويات: ", "خصم أو تسوية مباشرة بتاريخ ووصف يظهران في كشف الحساب."),
            ("كشف الحساب: ", "رصيد افتتاحي، فواتير، دفعات، خصومات، رصيد جارٍ بعد كل حركة، ورصيد ختامي قابل للطباعة."),
        ],
        compact=True,
    )

    add_page_break(doc)
    add_heading(doc, "8. الإعدادات والكتالوجات", 1)
    add_body(doc, "الإدارة تتحكم بالبيانات التي يتكرر استخدامها بدل كتابتها من جديد في كل حالة.")
    add_list(
        doc,
        [
            ("المواد: ", "الاسم والسعر وأنواع الأعمال المرتبطة بها، والمراحل المطلوبة، وهل تُحسب كوحدة في التقارير."),
            ("أنواع الأعمال: ", "تعريف Job Types وتحديد إذا كان الحساب على الأسنان أو الفك وربطها بالمواد المناسبة."),
            ("الزرعات: ", "كتالوج Implants يُستخدم داخل الحالة وفي التقارير."),
            ("التنظيم والجودة: ", "إدارة Tags وأسباب الفشل لتوحيد التصنيف والتقارير."),
        ],
    )

    add_heading(doc, "9. المستخدمون والصلاحيات", 1)
    add_body(doc, "يمكن إنشاء وتعديل وتعطيل حسابات الموظفين، ثم إعطاء كل شخص ما يحتاجه حسب دوره: الحالات، مراحل الإنتاج، التوصيل، العملاء، الحسابات، التقارير، والجودة. الحساب المعطّل لا يستطيع الدخول.")
    add_list(
        doc,
        [
            "الصلاحيات تفصل بين العرض، الإضافة، التعديل، الحذف، والتنفيذ في الأجزاء الحساسة.",
            "يوجد دعم لتسجيل وقت الدخول وعنوان الشبكة والمتصفح عندما تكون جداول التدقيق مفعّلة في بيئة المختبر.",
            "النظام يدعم العربية والإنجليزية واتجاهي RTL وLTR، لكن بعض النصوص ما زالت إنجليزية، لذلك الترجمة ليست كاملة 100%. ",
        ],
        compact=True,
    )

    add_heading(doc, "10. أكثر من مختبر", 1)
    add_body(doc, "بنية المنصة تسمح بتشغيل مختبرات منفصلة، لكل واحد دومين وقاعدة بيانات وعملة وهوية خاصة. مدير المنصة يجهز المختبر ويُنشئ أول حساب مدير ويتابع حالة التجهيز. هذه إدارة مركزية للمنصة، وليست شاشة يومية لموظف المختبر.")

    add_page_break(doc)
    add_heading(doc, "11. لوحات المتابعة والتقارير", 1)
    add_body(doc, "الهدف من التقارير هو قراءة الشغل من أكثر من زاوية: من ينتج، ماذا ننتج، لمن، وبأي جودة وقيمة.")
    add_list(
        doc,
        [
            "عدد الوحدات حسب الطبيب والمادة.",
            "Job Types بالحالات أو الوحدات، وتقارير المواد والزرعات وAbutments المسجلة داخل الحالات.",
            "تقارير QC والأخطاء وأسبابها، وReject وRepeat وModify وRedo وSuccessful كعدد أو نسبة.",
            "تقارير الحالات والمادة والقيمة المفوترة وتاريخ التسليم، مع فلاتر زمنية واختيار عميل والطباعة في أغلب الشاشات.",
            "لوحة الإدارة للأرقام العامة، وOperations Dashboard وLive Monitor للعمل الجاري، وجدول التسليم للمواعيد.",
        ],
        compact=True,
    )
    add_callout(
        doc,
        "ملاحظة عن لوحة الإدارة",
        "إعداد Sample Data موجود ومفعّل افتراضياً في الكود. قبل اعتماد أرقام اللوحة كمؤشرات حقيقية يجب تعطيله والتحقق من البيئة الفعلية لكل مختبر.",
        fill=AMBER_LIGHT,
        title_color="9A6700",
    )

    add_heading(doc, "12. ما هو جاهز وما يحتاج استكمال", 1)
    add_body(doc, "حتى تكون الصورة صريحة، هذه الأجزاء موجودة كأساس أو كود قديم لكنها ليست خدمة مكتملة ننصح بوعد العميل بها الآن:")
    add_list(
        doc,
        [
            ("Client Portal: ", "تحتاج إكمال الدخول والصلاحيات قبل فتحها للعملاء."),
            ("تطبيق الهاتف وواجهات API القديمة: ", "تحتاج مراجعة أمنية وتحديث طريقة المصادقة واختبار المسارات."),
            ("الإشعارات: ", "بعض الأكواد موجودة، لكن تحتاج إعداد FCM واختبار فعلي قبل اعتبارها ميزة عاملة."),
            ("Devices وBuilds: ", "ليست مفعّلة حالياً، ولا يوجد اعتماد عليها في التشغيل الحالي."),
            ("External Labs وتوصيل Abutments: ", "الواجهات أو الأكواد موجودة جزئياً لكن المسارات غير مفعّلة."),
            ("إدارة المخزون: ", "غير موجودة كميزة كاملة حالياً، وهي المقترح التالي."),
        ],
        compact=True,
    )

    add_page_break(doc)
    add_heading(doc, "13. مقترح إضافة مخزون خفيف", 1)
    add_callout(
        doc,
        "الفكرة الأساسية",
        "نضيف مخزوناً يخدم عمل المختبر، مش نظام ERP كبير. التركيز يكون على الكمية الحالية، الاستلام، الصرف الفعلي، التسويات، التنبيه، وسجل الحركة.",
        fill=BLUE_LIGHT,
        title_color="1D4ED8",
    )
    add_heading(doc, "لماذا لا نخصم عند إنشاء الحالة؟", 2)
    add_body(
        doc,
        "لأن الحالة تخبرنا ماذا نريد أن نصنع، لكنها لا تخبرنا بدقة ماذا استُهلك. قرص الزركون الواحد قد ينتج ثمانية أسنان تقريباً، أو أكثر، أو أقل حسب المقاسات وترتيب الـNesting والكسر والبقايا. الخصم التلقائي عند إنشاء الحالة سيعطي أرقاماً تبدو دقيقة لكنها عملياً غير صحيحة.",
        keep=True,
    )

    add_heading(doc, "طريقة الاستهلاك المقترحة", 2)
    add_list(
        doc,
        [
            ("أقراص الزركون: ", "يُخصم قرص واحد عندما يفتح موظف Milling قرصاً جديداً، ثم ترتبط به الحالات أو الأعمال التي تم ترتيبها عليه. عند النهاية يُعلّم كمستهلك أو تالف، بدون افتراض عدد أسنان ثابت."),
            ("القطع المحددة: ", "مثل Abutments أو Screws؛ تُخصم بالقطعة عند صرفها للحالة."),
            ("المواد السائبة: ", "مثل Resin أو Plaster؛ يسجل القسم الكمية الفعلية على دفعة العمل أو نهاية الوردية حسب أسهل إجراء للمختبر."),
            ("التغليف: ", "الأكياس أو العلب تُخصم عند تجهيز الحالة للتسليم."),
            ("التصحيح: ", "أي فرق جرد أو تلف يدخل كتسوية بسبب واضح واسم المستخدم، فلا يتغير الرصيد بصمت."),
        ],
    )

    add_heading(doc, "مكونات المرحلة الأولى", 2)
    add_list(
        doc,
        [
            "بطاقة صنف: الاسم، الفئة، الوحدة، الكمية الحالية، الحد الأدنى، والحالة.",
            "أربع حركات أساسية: استلام، صرف، تسوية، وتلف، مع التاريخ والمستخدم والسبب.",
            "شاشة مواد منخفضة، سجل حركة لكل صنف، واستهلاك حسب المادة والفترة أو القسم.",
            "ربط اختياري بالحالة أو بمرحلة الإنتاج عندما يكون الربط مفيداً، بدون إجبار كل حركة على وجود حالة.",
            "صلاحيات منفصلة لمن يسجل الاستلام أو الصرف أو التسوية أو يشاهد التقارير.",
        ],
        compact=True,
    )

    add_heading(doc, "ما لا ندخله في البداية", 2)
    add_body(doc, "طلبات شراء متقدمة، مقارنة موردين، مستودعات متعددة، توقعات شراء آلية، ومحاسبة تكلفة صناعية كاملة. هذه ممكن تأتي لاحقاً إذا أثبت الاستخدام اليومي أن المختبر يحتاجها.")

    add_page_break(doc)
    add_heading(doc, "14. شكل التنفيذ المقترح", 1)
    add_list(
        doc,
        [
            ("الخطوة الأولى — الأساس: ", "الأصناف، الأرصدة الافتتاحية، الاستلام، الصرف، التسوية، التنبيه، والصلاحيات."),
            ("الخطوة الثانية — الربط العملي: ", "فتح قرص الزركون من Milling وربطه بالأعمال، وصرف القطع للحالة، والتغليف عند التسليم."),
            ("الخطوة الثالثة — القياس: ", "تقارير الاستهلاك والفروقات ومراجعة الإجراءات مع المستخدمين بعد فترة تشغيل حقيقية."),
            ("الخطوة الرابعة — التوسعة عند الحاجة: ", "إضافة المشتريات أو الموردين أو أكثر من موقع فقط إذا ظهر احتياج واضح."),
        ],
        numbered=True,
    )

    add_heading(doc, "أسئلة تساعد في اتخاذ القرار", 1)
    add_list(
        doc,
        [
            "ما هي أهم 20–30 مادة نريد ضبطها أولاً بدل إدخال كل شيء من اليوم الأول؟",
            "من يسجل الاستلام، ومن يملك صلاحية الصرف أو التسوية؟",
            "هل تسجيل استهلاك المواد السائبة سيكون لكل Build أم مرة في نهاية الوردية؟",
            "هل نحتاج تنبيه داخل النظام فقط، أم أيضاً رسالة أو بريد عند انخفاض الكمية؟",
            "ما التقارير التي ستغيّر قراراً فعلياً: النواقص، الاستهلاك، التلف، أم التكلفة؟",
        ],
    )

    add_callout(
        doc,
        "الخلاصة",
        "Solent يغطي اليوم قلب عمل المختبر: العملاء، الحالات، الإنتاج، الجودة، التوصيل، الحسابات والتقارير. أفضل توسعة للمخزون هي توسعة صغيرة مبنية على الاستهلاك الحقيقي داخل المرحلة، وليس خصماً نظرياً عند إنشاء الحالة. بهذا يبقى النظام قريباً من حجمه الأصلي ويعطي أرقاماً يمكن الوثوق بها.",
    )

    # Normalize paragraphs and remove accidental empty trailing content.
    for paragraph in doc.paragraphs:
        set_paragraph_rtl(paragraph, paragraph.alignment or WD_ALIGN_PARAGRAPH.RIGHT)
        for run in paragraph.runs:
            set_run_font(run)

    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
